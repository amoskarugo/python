from docx import Document
from docx.shared import Inches

document = Document()
# profile picture
document.add_picture("me.jpg", width=Inches(2.0))

# email, phone and name
name = input("Enter your name: ")
phone_number = input("Enter your phone number: ")
email = input("Enter your email address: ")

document.add_paragraph(name + " | " + phone_number + " | " + email)

# about me
document.add_heading("ABOUT ME")
document.add_paragraph(input("Tell me about yourself? "))
document.save("cv.docx")

